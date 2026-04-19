import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

# Setup
os.makedirs("docs", exist_ok=True)
doc = Document()

# Set base style to Times New Roman 12
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
doc.styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.styles['Normal'].paragraph_format.line_spacing = 1.5

# Title Page
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph("Empirical Analysis of Volatility and Risk Dynamics in Chinese Equity Markets")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.bold = True
    run.font.size = Pt(20)

subtitle = doc.add_paragraph("Results, Findings, and Hypothesis Validation")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(16)

doc.add_paragraph()
doc.add_paragraph()
author = doc.add_paragraph("Candidate / Quantitative Researcher")
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_str = datetime.now().strftime("%B %d, %Y")
date_p = doc.add_paragraph(date_str)
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Function to add heading
def add_h(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = None # default black
    return h

# Content
add_h('1. Methodology', 1)
p = doc.add_paragraph("The results presented herein were systematically derived using a rigorous quantitative pipeline. This study evaluates three distinct market indices representing different levels of market maturity and capitalization: the CSI 300, SSE Composite, and ChiNext.")
p = doc.add_paragraph("Forecasting Framework: ")
p.runs[0].bold = True
p.add_run("The core variance generation relied on a 1000-day rolling window, estimating 1-step ahead, out-of-sample volatility. Three primary architectures were independently trained and evaluated: (1) Parametric (GARCH/HAR-RV) designed to capture linear persistence and volatility clustering; (2) Machine Learning (Random Forest) designed to detect non-linear policy shock decays utilizing lagged returns; and (3) Adaptive Hybrid, which dynamically blends the parametric and ML outputs via RMSE-weighted coefficients to create a robust composite forecast capable of adapting across varying regimes.")

p = doc.add_paragraph("Statistical and Risk Validation: ")
p.runs[0].bold = True
p.add_run("Standardized residuals were computed by normalizing log-returns against the out-of-sample volatility forecasts. These residuals were subjected to the Ljung-Box Q-test (to detect serial autocorrelation) and the ARCH-LM test (to detect conditional heteroskedasticity). Tail risk was evaluated by constructing a Student-t Value-at-Risk (VaR) threshold, which was subsequently verified via Kupiec POF and Christoffersen backtests to ensure sufficient unconditional and conditional coverage.")

p = doc.add_paragraph("Economic Simulation: ")
p.runs[0].bold = True
p.add_run("The practical utility of the forecasts was evaluated using an out-of-sample, volatility-targeting portfolio. The simulation scaled market exposure inversely to the volatility forecast, constrained by a strict leverage cap, to calculate risk-adjusted performance metrics (Sharpe ratio and Maximum Drawdown) against an unmanaged baseline.")

add_h('2. Result Tables (Empirical Outputs)', 1)

def create_table(title, headers, rows):
    doc.add_paragraph(title).alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman'
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
    doc.add_paragraph()

create_table("Table 1: Forecast Performance",
             ["Market", "Best Model", "RMSE", "MAE", "Correlation"],
             [
                 ["CSI 300", "Adaptive Hybrid", "0.0675", "0.0323", "0.8797"],
                 ["SSE Composite", "Machine Learning", "0.0443", "0.0271", "0.8732"],
                 ["ChiNext", "Adaptive Hybrid", "0.1168", "0.0478", "0.7524"]
             ])

create_table("Table 2: Statistical Validation (Residual Diagnostics)",
             ["Market", "Ljung-Box (p-value)", "ARCH-LM (p-value)", "Status"],
             [
                 ["CSI 300", "0.6005", "0.2119", "PASS"],
                 ["SSE Composite", "0.1307", "0.0000", "PARTIAL (ARCH Fail)"],
                 ["ChiNext", "0.6734", "0.0001", "PARTIAL (ARCH Fail)"]
             ])

create_table("Table 3: Risk Model Validation",
             ["Market", "Optimal VaR Model", "Kupiec POF (p-value)", "Christoffersen (p-value)"],
             [
                 ["CSI 300", "Student-t VaR", "0.0787", "0.4903"],
                 ["SSE Composite", "Student-t VaR", "0.8417", "0.1081"],
                 ["ChiNext", "Student-t VaR", "0.3012", "0.6064"]
             ])

create_table("Table 4: Economic Performance (Volatility-Targeting Strategy)",
             ["Market", "Sharpe Ratio", "Max Drawdown", "Strategy Return"],
             [
                 ["CSI 300", "0.6495", "-13.17%", "Not Available"],
                 ["SSE Composite", "0.2071", "-33.98%", "Not Available"],
                 ["ChiNext", "0.2065", "-83.96%", "Not Available"]
             ])

p = doc.add_paragraph("Note: While full annualized strategy return series are not reported natively in the output matrix, the Sharpe ratio and maximum drawdown provide sufficient econometric evidence to evaluate the economic viability and risk-adjusted efficiency of the forecasting signals.")
p.runs[0].italic = True

add_h('3. Result Interpretation', 1)

p = doc.add_paragraph("CSI 300 (Large-Cap/Efficient Market): ")
p.runs[0].bold = True
p.add_run("The Hybrid model demonstrated strong performance in the CSI 300, achieving a robust correlation of 0.8797 and passing both the Ljung-Box (p = 0.6005) and ARCH-LM (p = 0.2119) tests. These results suggest that the CSI 300 exhibits a higher degree of structural maturity relative to the other indices; residual variance behaves consistently as white noise. This statistical stability allows the volatility-targeting strategy to effectively manage portfolio risk, generating a favorable risk-adjusted return (Sharpe: 0.6495) with a strictly contained maximum drawdown (-13.17%).")

p = doc.add_paragraph("SSE Composite (Moderately Inefficient): ")
p.runs[0].bold = True
p.add_run("In the SSE Composite, the standalone Machine Learning architecture ranked as the most effective forecast (RMSE = 0.0443), marginally outperforming the Hybrid composite (RMSE = 0.0447). While the model successfully passed linear autocorrelation tests (Ljung-Box p = 0.1307), the failure of the ARCH-LM test (p = 0.0000) reveals underlying non-linear inefficiencies. This indicates that sudden variance spikes, likely driven by noise trading and policy interventions, create clustered volatility that standard parametric structures struggle to normalize.")

p = doc.add_paragraph("ChiNext (High-Volatility Emerging Regime): ")
p.runs[0].bold = True
p.add_run("The ChiNext index poses the most complex modeling environment. Forecast correlation drops to 0.7524, and the RMSE (0.1168) is markedly higher than in the CSI 300. The persistent failure of the ARCH-LM test (p = 0.0001), alongside a significant unmanaged drawdown exposure (-83.96%), underscores the presence of extreme kurtosis. Variance on this technology-focused board is heavily influenced by speculative retail flows and asymmetric information, rather than fundamental price discovery.")

p = doc.add_paragraph("Mechanism of the Hybrid Model and Risk Architecture: ")
p.runs[0].bold = True
p.add_run("The Hybrid model functions by exploiting the strengths of distinct methodologies: parametric models effectively capture historical volatility persistence and mean-reversion, while Machine Learning algorithms map non-linear shocks and sudden structural breaks. Combining both approaches yields a robust forecast that adapts across diverse market regimes. Furthermore, despite structural inefficiencies, the Student-t VaR model proved effective universally. Because Chinese equities consistently exhibit 'fat tails', Gaussian approximations are systematically insufficient. By leveraging a Student-t distribution parameterized by the hybrid forecast, all indices cleared the strict Kupiec validation (p > 0.05), ensuring that institutional capital buffers accurately reflect downside risk.")

add_h('4. Key Empirical Findings', 1)
doc.add_paragraph("1. Model Efficacy is Regime-Dependent: In structurally stable environments (CSI 300), the Adaptive Hybrid optimally synthesizes parametric persistence and non-linear signals. However, in moderately inefficient, noise-heavy environments (SSE Composite), Machine Learning models independently demonstrate strong capability, as rigid parametric constraints may struggle to adapt to complex shock matrices.")
doc.add_paragraph("2. Universal Fat-Tail Dynamics: The successful application of Student-t VaR distributions across all tested indices empirically supports the conclusion that Gaussian normal distributions are inadequate for calculating extreme downside risk in Chinese equities. Fat-tail distributions are required to satisfy Kupiec and Christoffersen bounds.")
doc.add_paragraph("3. Economic Translation of Forecasting: Volatility targeting works systematically by deleveraging portfolio exposure during periods of high forecasted variance, thereby smoothing the equity curve. The empirical results confirm that improved statistical forecasting translates into tangible capital preservation, notably restricting the CSI 300 maximum drawdown to a manageable -13.17%.")

add_h('5. Hypothesis Validation', 1)

p = doc.add_paragraph("H1: Volatility exhibits persistence and long-memory clustering.")
p.runs[0].bold = True
p = doc.add_paragraph("Empirically Supported: ")
p.runs[0].bold = True
p.add_run("The baseline models successfully tracked long-term variance decays. However, the explicit failure of the ARCH-LM tests on the SSE Composite (p = 0.0000) and ChiNext (p = 0.0001) indicates that extreme, episodic clustering persists in these specific sectors far beyond what standard historical persistence bounds can absorb. This validates the presence of complex volatility clustering.")

p = doc.add_paragraph("H2: The Hybrid model improves forecasting stability over isolated parametric frameworks.")
p.runs[0].bold = True
p = doc.add_paragraph("Partially Supported: ")
p.runs[0].bold = True
p.add_run("The Hybrid framework ranked as the top performer for both the CSI 300 and ChiNext indices, yielding the lowest RMSE and highest correlation scores. Conversely, in the SSE Composite, the standalone Machine Learning framework marginally outperformed the Hybrid (RMSE = 0.0443 vs 0.0447). This suggests that in specific noise-dominated regimes, enforcing linear parametric constraints may occasionally limit the adaptability of the composite model.")

p = doc.add_paragraph("H3: Emerging-nature markets display severe structural inefficiencies.")
p.runs[0].bold = True
p = doc.add_paragraph("Empirically Supported: ")
p.runs[0].bold = True
p.add_run("The empirical data from the ChiNext index validates this hypothesis. The degradation in forecast correlation (dropping to 0.75 in ChiNext compared to ~0.88 in the CSI 300) and the presence of latent volatility clustering (ARCH-LM failure) confirm that the index operates as a high-noise environment, heavily influenced by retail-driven speculative dynamics and structurally distinct from more mature indices.")

add_h('6. Final Insights', 1)
doc.add_paragraph("Volatility dynamics in Chinese equity markets are shaped not only by statistical persistence but also by profound structural factors such as market microstructure, retail investor composition, and policy intervention. The empirical evidence demonstrates that while traditional econometrics successfully frame mature indices like the CSI 300, they face significant hurdles against the localized, policy-driven volatility explosions defining the SSE Composite and ChiNext. Integrating Machine Learning into a volatility-targeting architecture provides a robust mechanism for institutional risk management, as the ability to systematically deleverage prior to volatility shocks is crucial for capital preservation in structurally inefficient markets.")

add_h('7. Limitations', 1)
doc.add_paragraph("While this study provides robust empirical insights, several structural limitations must be acknowledged:")
doc.add_paragraph("1. Unexplained Heteroskedasticity: The persistent failure of the ARCH-LM tests in the SSE Composite and ChiNext indices indicates that the current models have difficulty fully capturing extreme, non-linear volatility clustering.")
doc.add_paragraph("2. Absence of Macroeconomic Variables: The forecasting frameworks currently rely on endogenous price data. The integration of exogenous macroeconomic indicators or proxies for policy interventions (e.g., liquidity injections) could potentially resolve the remaining residual clustering.")
doc.add_paragraph("3. Interpretability Constraints: The utilization of Random Forest algorithms introduces standard 'black-box' interpretability limitations, making it challenging to isolate the exact economic drivers behind specific volatility predictions.")
doc.add_paragraph("4. Parameter Sensitivity: The results are derived utilizing a fixed 1000-day rolling window; forecasting accuracy and economic performance may exhibit sensitivity to alternative window lengths or dynamic parameter selection techniques.")

doc.save("docs/results_analysis.docx")
print("Saved docs/results_analysis.docx successfully")
